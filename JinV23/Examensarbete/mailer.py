from bs4 import BeautifulSoup
from datetime import datetime
from docx import Document
import logging
import markdown2
import os
import pandas as pd
import random
import re
import win32com.client as win32

# Debug flagga
debug = False

# Skapa en loggfil med datum och tid i namnet
log_filename = f'logs/mailer_{datetime.now().strftime("%Y-%m-%d")}.log'

# empty the log file if it exists
open(log_filename, 'w').close()

# Skapa logs mappen om den inte finns
if not os.path.exists('logs'):
    os.makedirs('logs')

# Konfigurera loggning
logging.basicConfig(filename=log_filename,
                    level=logging.DEBUG,
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                    encoding='utf-8')

logger = logging.getLogger(__name__)

emails_created = 0
words_created = 0
files_found = 0
errors_encountered = 0

def anonymize_data(row):
    return {
        'Förnamn': 'student',
        'Efternamn': 'lastname',
        'E-post': 'email',
        'Användarnamn': 'username',
        'Gafe-epostadress': 'gafeemail'
    }

def extract_email(text):
    try:
        # Check for the [name](mailto:email) format
        match = re.search(r'\[.*?\]\(mailto:(.*?)\)', text)
        if match:
            return match.group(1)
        # Check for the <email> format
        match = re.search(r'<(.*?)>', text)
        if match:
            return match.group(1)
        # Otherwise, assume it's a plain email address
        return text
    except Exception as e:
        logger.error(f"Error extracting email from text '{text}': {e}")
        return ""

def read_markdown_file(filepath):
    try:
        with open(filepath, 'r', encoding='utf-8') as file:
            content = ''
            for line in file:
                if line.strip() == '---':
                    break
                content += line
        return content
    except Exception as e:
        logger.error(f"Fel vid läsning av fil {filepath}: {e}")
        raise

def handle_html_content(content):
    try:
        if "Detaljer om projektet:" in content:
            parts = content.split("Detaljer om projektet:")
            if len(parts) > 1:
                content = parts[0] + "Detaljer om projektet:" + convert_to_bullet_list(parts[1])
            else:
                logger.warning("No content found after 'Detaljer om projektet:'")
        return content
    except Exception as e:
        logger.error(f"Error handling HTML content: {e}", exc_info=True)
        return content

def convert_to_bullet_list(text):
    try:
        lines = text.split('\n')
        bullet_items = [line.strip() for line in lines if line.strip() and not line.strip() == '-']
        return '<ul>' + ''.join(f'<li>{line}</li>' for line in bullet_items) + '</ul>'
    except Exception as e:
        logger.error(f"Error converting text to bullet list: {e}")
        return text

def process_html_conversion(content):
    try:
        return markdown2.markdown(content)
    except Exception as e:
        logger.error(f"Fel vid konvertering av markdown till HTML: {e}")
        raise

def prepare_personal_greeting(firstname):
    return f"<p>Hej {firstname},</p><p>här kommer den efterlängtade feedbacken på ditt projekt. <br>Jag ber hemskt mycket om ursäkt att det dröjt så länge, <br>men bättre sent än ännu senare.</p>"

def prepare_closure():
    return "<p><br>Lycka till med LIAn<br>och ha en jätteskön sommar!<br>Varma hälsningar,<br>Marcus</p>"

def create_email(outlook, recipients, html_content, debug):
    try:
        mail = outlook.CreateItem(0)
        mail.To = recipients if not debug else 'marmed02@gafe.molndal.se'
        mail.Subject = 'Feedback på examensarbete'
        mail.HTMLBody = html_content
        mail.Save()
        logger.info(f"E-postmeddelande skapat och sparat som utkast för {recipients}")
        return True
    except Exception as e:
        logger.error(f"Fel vid skapande av e-postmeddelande för {recipients}: {e}")
        return False

def create_word_document(filepath, firstname, lastname, userid, html_content):
    try:
        doc = Document()
        doc.add_heading('Feedback på Examensarbete', 0)
        doc.add_paragraph(f'Namn: {firstname} {lastname} ({userid})')
        doc.add_heading('Innehåll', level=1)
        add_html_to_docx(doc, html_content)
        doc.save(filepath)
        logger.info(f"Word-dokument skapat och sparat för {firstname} {lastname}")
        return True
    except Exception as e:
        logger.error(f"Fel vid skapande av Word-dokument för {firstname} {lastname}: {e}")
        return False

def add_html_to_docx(doc, html):
    try:
        soup = BeautifulSoup(html, 'html.parser')
        for element in soup.descendants:
            if element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name == 'br':
                doc.add_paragraph()
            elif element.name == 'h1':
                doc.add_heading(element.text, level=1)
            elif element.name == 'h2':
                doc.add_heading(element.text, level=2)
            elif element.name == 'h3':
                doc.add_heading(element.text, level=3)
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Bullet')
            elif element.name == 'ol':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Number')
    except Exception as e:
        logger.error(f"Error adding HTML to Word document: {e}")

def process_student(row, outlook, debug):
    global emails_created, words_created, files_found, errors_encountered
    try:
        epost = extract_email(row['E-post'])
        gafe_epost = extract_email(row['Gafe-epostadress'])

        filnamn = f'{row["Användarnamn"].strip()}.md'
        if os.path.isfile(filnamn):
            files_found += 1
            logger.debug(f"Fil hittad för {row['Förnamn']} {row['Efternamn']} ({row['Användarnamn']})")
            content = read_markdown_file(filnamn)
            content = handle_html_content(content)
            html_content = process_html_conversion(content)
            
            personal_greeting = prepare_personal_greeting(row['Förnamn'])
            closure = prepare_closure()
            full_html_content = personal_greeting + html_content + closure

            email_success = create_email(outlook, f'{epost}; {gafe_epost}', full_html_content, debug)
            word_success = create_word_document(
                f'feedback/{row["Förnamn"].strip().capitalize()} {row["Efternamn"].strip().capitalize()} ({row["Användarnamn"].strip()}).docx', 
                row['Förnamn'].strip().capitalize(), 
                row['Efternamn'].strip().capitalize(), 
                row['Användarnamn'].strip(), 
                full_html_content
            )

            if email_success:
                emails_created += 1
            if word_success:
                words_created += 1

        else:
            logger.warning(f"Fil {filnamn} för {row['Förnamn']} {row['Efternamn']} ({row['Användarnamn']}) existerar inte")

    except KeyError as e:
        errors_encountered += 1
        logger.error(f"Fel vid bearbetning av student {row['Förnamn']} {row['Efternamn']}: kolumnen '{e}' saknas i DataFrame")
    except Exception as e:
        errors_encountered += 1
        logger.error(f"Fel vid bearbetning av student {row['Förnamn']} {row['Efternamn']}: {e}", exc_info=True)

def read_and_process_csv_file(file_path):
    try:
        # Read the CSV file into a DataFrame
        df = pd.read_csv(file_path, delimiter=',')
        logger.info("Läst CSV-filen till DataFrame")

        logger.debug(f"Initial columns in DataFrame: {df.columns.tolist()}")

        # Define expected column headers
        expected_columns = ['Förnamn', 'Efternamn', 'E-post', 'Användarnamn', 'Gafe-epostadress']
        
        # Check for missing columns
        missing_columns = [col for col in expected_columns if col not in df.columns]
        if missing_columns:
            raise ValueError(f"Missing required columns in DataFrame: {missing_columns}")

        logger.debug(f"Processed DataFrame: \n{df.head()}")
        return df
    except Exception as e:
        logger.error(f"Error reading and processing CSV file: {e}", exc_info=True)
        raise

def initialize_outlook():
    try:
        outlook = win32.Dispatch('outlook.application')
        logger.info("Startat Outlook Application")
        return outlook
    except Exception as e:
        logger.error(f"Error initializing Outlook application: {e}")
        raise

def create_feedback_folder():
    try:
        if not os.path.exists('feedback'):
            os.makedirs('feedback')
        logger.info("Skapat mappen 'feedback'")
    except Exception as e:
        logger.error(f"Error creating feedback folder: {e}")
        raise

def print_and_log_report():
    global emails_created, words_created, files_found, errors_encountered
    try:
        success = (emails_created == words_created == files_found) and (errors_encountered == 0)
        logger.info("Rapport:")
        logger.info(f"Skapade e-postmeddelanden: {emails_created}")
        logger.info(f"Skapade Word-dokument: {words_created}")
        logger.info(f"Hittade studentfiler: {files_found}")
        logger.info(f"Antal fel: {errors_encountered}")
        logger.info(f"Totalt lyckades: {success}")

        print("Rapport:")
        print(f"Skapade e-postmeddelanden: {emails_created}")
        print(f"Skapade Word-dokument: {words_created}")
        print(f"Hittade studentfiler: {files_found}")
        print(f"Antal fel: {errors_encountered}")
        print(f"Totalt lyckades: {success}")
    except Exception as e:
        logger.error(f"Error printing and logging report: {e}")
        raise

def start_outlook():
    try:
        outlook = win32.Dispatch('outlook.application')
        logger.info("Startat Outlook Application")
        return outlook
    except Exception as e:
        logger.error(f"Error starting Outlook application: {e}", exc_info=True)
        raise

def create_feedback_folder():
    try:
        if not os.path.exists('feedback'):
            os.makedirs('feedback')
        logger.info("Skapat mappen 'feedback'")
    except Exception as e:
        logger.error(f"Error creating feedback folder: {e}", exc_info=True)
        raise

def main():
    global emails_created, words_created, files_found, errors_encountered
    emails_created = 0
    words_created = 0
    files_found = 0
    errors_encountered = 0

    try:
        df = read_and_process_csv_file('emails.csv')
        outlook = start_outlook()
        create_feedback_folder()

        logger.info("Startar mailspammandet mwahahahaaaaa...")
        for index, row in df.iterrows():
            try:
                process_student(row, outlook, debug=False)
            except Exception as e:
                logger.error(f"Error processing student at index {index}: {e}", exc_info=True)
                errors_encountered += 1

        success = (emails_created == words_created == files_found) and (errors_encountered == 0)
        logger.info("Rapport:")
        logger.info(f"Skapade e-postmeddelanden: {emails_created}")
        logger.info(f"Skapade Word-dokument: {words_created}")
        logger.info(f"Hittade studentfiler: {files_found}")
        logger.info(f"Antal fel: {errors_encountered}")
        logger.info(f"Totalt lyckades: {success}")

        print("Rapport:")
        print(f"Skapade e-postmeddelanden: {emails_created}")
        print(f"Skapade Word-dokument: {words_created}")
        print(f"Hittade studentfiler: {files_found}")
        print(f"Antal fel: {errors_encountered}")
        print(f"Totalt lyckades: {success}")

    except Exception as e:
        logger.error(f"Ett fel inträffade: {e}", exc_info=True)

if __name__ == "__main__":
    main()
